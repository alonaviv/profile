from functools import wraps
from io import BytesIO
from urllib.parse import quote

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import render, reverse, redirect
from django.template.loader import render_to_string
from hebrew_numbers import int_to_gematria
from weasyprint import HTML, CSS
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from profile_server.pronouns import PronounWordDictionary, PronounOptions
from utils.date_helpers import get_printable_date, TrimesterType
from utils.school_dates import get_current_trimester
from .models import Evaluation, Class, Student, House

"""
Every view in the site needs to pass the teacher object in the context, so it can display the name
of the logged in teacher in the navbar. Note that teacher in this context is the TeacherUser, not the Teacher object which is used to verify that the user is expected.
"""


def main_evaluations_page(request):
    context = {'teacher': request.user}
    return render(request, 'evaluations/index.html', context)


def populate_evaluations_in_teachers_classes(teacher):
    """
    Make sure that there is an evaluation object in the DB for each student in each class of the given teacher.
    If there isn't, create an entry with empty text.
    """
    if not teacher.is_anonymous:
        current_trimester = get_current_trimester()

        for evaluated_class in teacher.class_set.filter(hebrew_year=current_trimester.hebrew_school_year):
            for student in evaluated_class.students.all():
                try:
                    Evaluation.objects.get(student=student, evaluated_class=evaluated_class,
                                           hebrew_year=current_trimester.hebrew_school_year,
                                           trimester=current_trimester.name)
                except Evaluation.DoesNotExist:
                    current_trimester = get_current_trimester()
                    evaluation = Evaluation(student=student, evaluated_class=evaluated_class,
                                            trimester=current_trimester.name,
                                            hebrew_year=current_trimester.hebrew_school_year)
                    evaluation.full_clean()
                    evaluation.save()


@login_required
def write_class_evaluations(request, class_id, anchor=None):
    current_trimester = get_current_trimester()

    teacher = request.user
    try:
        class_to_evaluate = Class.objects.get(id=class_id, hebrew_year=current_trimester.hebrew_school_year)
    except Class.DoesNotExist:
        return render(request, 'common/general_error_page.html',
                      {'error_message': 'השיעור המבוקש אינו קיים במערכת'})

    if class_to_evaluate.teacher != teacher:
        return render(request, 'common/general_error_page.html',
                      {'error_message': 'רק המורה של השיעור רשאי/ת לכתוב דיווחים'})

    if request.method == 'POST':
        evaluation = Evaluation.objects.get(id=request.POST['evaluation_id'])
        if 'evaluation_text' in request.POST:
            evaluation.evaluation_text = request.POST['evaluation_text']
        error_encountered = False

        if 'save_draft' in request.POST:  # Saves new text, while remaining unsubmitted
            if evaluation.is_submitted:
                return render(request, 'common/general_error_page.html',
                              {'error_message': 'לא ניתן לשמור טיוטא של דיווח שכבר נשלח'})

        elif 'submit' in request.POST:
            evaluation.is_submitted = True

            if evaluation.is_empty:
                messages.error(request, "לא ניתן להגיש דיווח ריק.", extra_tags=str(evaluation.id))
                error_encountered = True

        elif 'withdraw' in request.POST:
            evaluation.is_submitted = False

        else:
            return render(request, 'common/general_error_page.html',
                          {'error_message': 'תקלה בשליחת הדיווח'})

        if not error_encountered:
            evaluation.full_clean()
            evaluation.save()

    evaluations = Evaluation.objects.filter(evaluated_class=class_to_evaluate,
                                            trimester=current_trimester.name).order_by('student')

    if not anchor:
        anchor = f"anchor_{request.POST['evaluation_id']}" if 'evaluation_id' in request.POST else 'top'

    context = {
        'teacher': teacher, 'class': class_to_evaluate, 'evaluations': evaluations,
        'anchor': anchor
    }
    return render(request, 'evaluations/write_evaluations.html', context)


@login_required
def remove_evaluation(request, evaluation_id):
    teacher = request.user
    evaluation = Evaluation.objects.get(id=evaluation_id)

    if evaluation.is_submitted:
        messages.error(request, "לא ניתן לבטל דיווח שכבר נשלח לחונכ/ת", extra_tags=str(evaluation.id))

        return redirect(
            reverse('write_class_evaluations_with_anchor',
                    args=(evaluation.evaluated_class.id, f"anchor_{evaluation_id}")))

    if not evaluation.is_empty:
        messages.error(request,
                       "לא ניתן לבטל דיווח כאשר קיימת טיוטא. אם ברצונך לבטל את הדיווח בכל זאת, יש לשמור טיוטא ריקה.",
                       extra_tags=str(evaluation.id))

        return redirect(
            reverse('write_class_evaluations_with_anchor',
                    args=(evaluation.evaluated_class.id, f"anchor_{evaluation_id}")))

    if teacher != evaluation.evaluated_class.teacher:
        return render(request, 'common/general_error_page.html',
                      {'error_message': 'רק המורה של השיעור רשאי/ת להסיר דיווחים של השיעור'})
    if evaluation.is_student_in_class:
        return render(request, 'common/general_error_page.html',
                      {'error_message': 'לא ניתן להסיר דיווח של תלמיד/ה שעדיין נמצא/ת בשיעור'})

    evaluation.delete()

    return redirect(
        reverse('write_class_evaluations_with_anchor', args=(evaluation.evaluated_class.id, f"anchor_{evaluation_id}")))


@login_required
def write_evaluations_main_page(request):
    current_trimester = get_current_trimester()
    teacher = request.user

    populate_evaluations_in_teachers_classes(teacher)

    if teacher:
        classes = teacher.class_set.filter(hebrew_year=current_trimester.hebrew_school_year)
    else:
        classes = []

    classes = sorted(classes.all(), key=lambda klass: klass.name)

    context = {'classes': classes, 'teacher': teacher}
    return render(request, 'evaluations/write_evaluations_index.html', context)


def _get_student_evaluation_context(request, student_id):
    teacher = request.user

    if not teacher.is_homeroom_teacher:
        return redirect(reverse('not_homeroom_teacher_error'))

    try:
        student = Student.objects.get(id=student_id)
    except Student.DoesNotExist:
        return render(request, 'common/general_error_page.html',
                      {'error_message': 'תלמיד/ה זה/זו לא נמצא/ת בבית הספר'})

    if student.homeroom_teacher != teacher:
        return redirect(reverse('mismatched_homeroom_teacher_error'))

    pronoun_dict = PronounWordDictionary(teacher.pronoun_as_enum)

    return {'student': student, 'teacher': teacher, 'pronoun_dict': pronoun_dict,
            'trimester': get_current_trimester(), 'printable_date': get_printable_date()}


@login_required
def view_student_evaluations(request, student_id):
    context = _get_student_evaluation_context(request, student_id)

    if isinstance(context, HttpResponse):
        return context

    return render(request, 'evaluations/view_evaluations.html', context)


@login_required
def view_single_eval(request, evaluation_id):
    context = {'evaluation': Evaluation.objects.get(id=evaluation_id)}

    return render(request, 'evaluations/display_single_eval.html', context)


@login_required
def download_student_evaluations(request, student_id):
    context = _get_student_evaluation_context(request, student_id)

    if isinstance(context, HttpResponse):
        return context

    response = HttpResponse(content_type="application/pdf")

    trimester_name = context['trimester'].name.lower().replace("_"," ")
    trimester_year = context['trimester'].meeting_end_of_trimester.year
    filename = f"דיווחים של {context['student']} - {trimester_name} {trimester_year}"
    response['Content-Disposition'] = f"attachment; filename*=UTF-8''{quote(filename)}.pdf"
    html_string = render_to_string("evaluations/evaluations_as_pdf.html", context)
    with open('profile_server/static/css/style_for_pdf.css') as f:
        css_string = f.read()

    HTML(string=html_string).write_pdf(response, stylesheets=[CSS(string=css_string)])
    return response


@login_required
def view_evaluations_main_page(request):
    teacher = request.user

    if not teacher.is_homeroom_teacher:
        return redirect(reverse('not_homeroom_teacher_error'))

    if teacher:
        students = teacher.student_set.all()
    else:
        students = []

    students = sorted(students, key=lambda student: (student.first_name, student.last_name))
    context = {'students': students, 'teacher': teacher}
    return render(request, 'evaluations/view_evaluations_index.html', context)


@login_required
def evaluations_details(request, student_id):
    teacher = request.user
    current_trimester = get_current_trimester()

    if not teacher.is_homeroom_teacher:
        return redirect(reverse('not_homeroom_teacher_error'))

    student = Student.objects.get(id=student_id)

    evaluations = student.evaluations.filter(trimester=current_trimester.name,
                                             hebrew_year=current_trimester.hebrew_school_year).order_by(
        'evaluation_text')
    context = {
        'student': student, 'evaluations': evaluations, 'teacher': teacher
    }
    return render(request, 'evaluations/evaluation_details.html', context)


def failed_login(request):
    return render(request, 'evaluations/failed_login.html')


# ==== Historic evaluations download (admin-only) - any year/semester ===

# Only the three real trimesters are exposed - the null
# trimester is intentionally absent.
_HISTORIC_TRIMESTER_HEBREW_LABELS = {
    1: 'פגישה מספר 1 (ראשונה, מעט דיווחים)',
    2: 'פגישה מספר 2',
    3: 'פגישה מספר 3',
}


def _historic_staff_required(view_func):
    @wraps(view_func)
    @login_required
    def wrapped(request, *args, **kwargs):
        if not request.user.is_staff:
            return render(request, 'common/general_error_page.html',
                          {'error_message': 'רק משתמשים בעלי הרשאת ניהול רשאים לצפות בעמוד זה'})
        return view_func(request, *args, **kwargs)

    return wrapped


def _historic_hebrew_year_label(hebrew_year):
    """Same convention as Trimester.hebrew_school_year_printable (e.g. 'תשפ״ה')."""
    return int_to_gematria(hebrew_year - 5000)


class _HistoricTrimester:
    """
    Minimal stand-in for `utils.date_helpers.Trimester` that exposes only the
    attributes the PDF template reads. A real Trimester is built from school
    config (meeting dates / grace period) and can't represent arbitrary
    historical year+trimester pairs without that config, so we use this shim.
    """

    def __init__(self, trimester_number, hebrew_year):
        self.number = trimester_number
        self._hebrew_year = hebrew_year

    @property
    def hebrew_school_year_printable(self):
        return _historic_hebrew_year_label(self._hebrew_year)


class _HistoricStudent:
    """
    Read-only proxy around a real `Student` used only when rendering an historic
    PDF. The shared template `evaluations/evaluations_as_pdf.html` pulls its
    evaluations via:

        {% with all_student_evals=student.all_evals_in_current_trimester
                completed_student_evals=student.completed_evals_in_current_trimester %}

    so to render an arbitrary historical year/semester without modifying the
    template, we hand it a student-like object whose
    `*_evals_in_current_trimester` properties return the data for the requested
    historic year+trimester instead of the actual current trimester. Every other
    attribute (first_name, last_name, house, homeroom_teacher, ...) is delegated
    to the wrapped Student via `__getattr__`, so the rest of the template
    renders identically.
    """

    def __init__(self, student, hebrew_year, trimester_name):
        self._student = student
        self._hebrew_year = hebrew_year
        self._trimester_name = trimester_name

    def __getattr__(self, name):
        return getattr(self._student, name)

    def __str__(self):
        return str(self._student)

    @property
    def completed_evals_in_current_trimester(self):
        # Mirrors Student.completed_evals_in_current_trimester verbatim, but
        # for the requested historic year+trimester.
        completed_evals = []
        for evaluation in self._student.evaluations.filter(
                trimester=self._trimester_name,
                hebrew_year=self._hebrew_year,
                is_submitted=True):
            if evaluation.evaluation_text:
                completed_evals.append(evaluation)
        return completed_evals

    @property
    def all_evals_in_current_trimester(self):
        # Mirrors Student.all_evals_in_current_trimester verbatim, but for
        # the requested historic year+trimester.
        return list(self._student.evaluations.filter(
            trimester=self._trimester_name,
            hebrew_year=self._hebrew_year))


def _historic_homeroom_for_exports(student, hebrew_year):
    if hebrew_year != get_current_trimester().hebrew_school_year:
        return (
            PronounWordDictionary(PronounOptions.FEMALE),
            'דיווח היסטורי - אין חונך/ת נוכחי/ת',
        )

    homeroom_teacher = student.homeroom_teacher
    if homeroom_teacher is not None:
        return (
            PronounWordDictionary(homeroom_teacher.pronoun_as_enum),
            str(homeroom_teacher),
        )
    return PronounWordDictionary(PronounOptions.FEMALE), '— אין מחנך/ת נוכחי/ת —'


def _historic_submitted_evals_for_student(student):
    return student.evaluations.filter(is_submitted=True).exclude(evaluation_text='')


@_historic_staff_required
def historic_index(request):
    houses = House.objects.order_by('house_name')
    return render(request, 'evaluations/historic_index.html', {'houses': houses})


@_historic_staff_required
def historic_house(request, house_id):
    try:
        house = House.objects.get(id=house_id)
    except House.DoesNotExist:
        return render(request, 'common/general_error_page.html',
                      {'error_message': 'הבית המבוקש אינו קיים במערכת'})
    students = Student.objects.filter(house=house).order_by('first_name', 'last_name')
    return render(request, 'evaluations/historic_students.html',
                  {'house': house, 'students': students})


@_historic_staff_required
def historic_student(request, student_id):
    try:
        student = Student.objects.get(id=student_id)
    except Student.DoesNotExist:
        return render(request, 'common/general_error_page.html',
                      {'error_message': 'תלמיד/ה זה/זו לא נמצא/ת בבית הספר'})

    years = sorted(
        set(_historic_submitted_evals_for_student(student).values_list('hebrew_year', flat=True)),
        reverse=True,
    )
    year_entries = [{'hebrew_year': y, 'label': _historic_hebrew_year_label(y)} for y in years]
    return render(request, 'evaluations/historic_years.html',
                  {'student': student, 'year_entries': year_entries})


@_historic_staff_required
def historic_student_year(request, student_id, hebrew_year):
    try:
        student = Student.objects.get(id=student_id)
    except Student.DoesNotExist:
        return render(request, 'common/general_error_page.html',
                      {'error_message': 'תלמיד/ה זה/זו לא נמצא/ת בבית הספר'})

    trimester_names = set(
        _historic_submitted_evals_for_student(student)
        .filter(hebrew_year=hebrew_year)
        .values_list('trimester', flat=True)
    )

    semester_entries = []
    for trimester_name in trimester_names:
        try:
            number = TrimesterType[trimester_name].value
        except KeyError:
            continue
        if number in _HISTORIC_TRIMESTER_HEBREW_LABELS:
            semester_entries.append(
                {'number': number, 'label': _HISTORIC_TRIMESTER_HEBREW_LABELS[number]}
            )
    semester_entries.sort(key=lambda entry: entry['number'])

    context = {
        'student': student,
        'hebrew_year': hebrew_year,
        'hebrew_year_label': _historic_hebrew_year_label(hebrew_year),
        'semester_entries': semester_entries,
    }
    return render(request, 'evaluations/historic_semesters.html', context)


@_historic_staff_required
def historic_download(request, student_id, hebrew_year, trimester_num):
    try:
        student = Student.objects.get(id=student_id)
    except Student.DoesNotExist:
        return render(request, 'common/general_error_page.html',
                      {'error_message': 'תלמיד/ה זה/זו לא נמצא/ת בבית הספר'})

    if trimester_num not in _HISTORIC_TRIMESTER_HEBREW_LABELS:
        return render(request, 'common/general_error_page.html',
                      {'error_message': 'סמסטר לא תקין'})
                 
    trimester_name = TrimesterType(trimester_num).name

    pronoun_dict, teacher_for_template = _historic_homeroom_for_exports(student, hebrew_year)

    context = {
        # Wrap the student so the shared PDF template's
        # `{% with ... =student.completed_evals_in_current_trimester %}` block
        # surfaces the historical evals for the requested year+semester instead
        # of the current trimester's. Every other attribute is delegated to the
        # real Student via the proxy.
        'student': _HistoricStudent(student, hebrew_year, trimester_name),
        'teacher': teacher_for_template,
        'pronoun_dict': pronoun_dict,
        'trimester': _HistoricTrimester(trimester_num, hebrew_year),
        # The PDF template renders this string in its top-corner banner.
        # For the live teacher download it's just today's date - a "printed on"
        # stamp. For a historic download we additionally flag the document as
        # historical so a parent reading the PDF doesn't mistake it for a
        # fresh report.
        'printable_date': f'דיווח היסטורי - הופק {get_printable_date()}',
    }

    html_string = render_to_string('evaluations/evaluations_as_pdf.html', context)
    with open('profile_server/static/css/style_for_pdf.css') as f:
        css_string = f.read()

    response = HttpResponse(content_type='application/pdf')
    filename = (
        f"דיווחים היסטוריים של {student} - "
        f"שנת {_historic_hebrew_year_label(hebrew_year)}, פגישה {trimester_num}.pdf"
    )
    response['Content-Disposition'] = f"attachment; filename*=UTF-8''{quote(filename)}"

    HTML(string=html_string).write_pdf(response, stylesheets=[CSS(string=css_string)])
    return response


def _historic_build_evaluations_docx(student, hebrew_year, trimester_num):
    """
    Professional-looking .docx: black text only, larger header block, full-width
    horizontal rules between sections. Same evaluation content as the historic PDF.
    """
    trimester_name = TrimesterType(trimester_num).name
    proxy = _HistoricStudent(student, hebrew_year, trimester_name)
    completed = proxy.completed_evals_in_current_trimester

    pronoun_dict, teacher_line = _historic_homeroom_for_exports(student, hebrew_year)

    doc = Document()
    # Document-level hint so Word applies bidirectional layout (not only right align).
    settings_el = doc.settings.element
    if settings_el.find(qn('w:bidiVisual')) is None:
        settings_el.append(OxmlElement('w:bidiVisual'))

    def _rtl_paragraph_marks(paragraph):
        """RTL paragraph: bidi + jc=start (Word: start = visually right for RTL)."""
        p_pr = paragraph._p.get_or_add_pPr()
        if p_pr.find(qn('w:bidi')) is None:
            p_pr.append(OxmlElement('w:bidi'))
        jc = p_pr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            p_pr.append(jc)
        jc.set(qn('w:val'), 'start')

    def _run_rtl(run):
        """Mark Hebrew runs as complex-script RTL so Word orders glyphs correctly."""
        r_pr = run._element.get_or_add_rPr()
        if r_pr.find(qn('w:rtl')) is None:
            r_pr.append(OxmlElement('w:rtl'))

    def _rtl_paragraph(text, size_pt=None, bold=False, space_before=0, space_after=8):
        p = doc.add_paragraph()
        _rtl_paragraph_marks(p)
        fmt = p.paragraph_format
        fmt.space_before = Pt(space_before)
        fmt.space_after = Pt(space_after)
        fmt.left_indent = Pt(0)
        fmt.right_indent = Pt(0)
        fmt.first_line_indent = Pt(0)
        run = p.add_run(text)
        _run_rtl(run)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        run.bold = bold
        return p

    def _horizontal_rule():
        """Full-width horizontal line (paragraph bottom border)."""
        p = doc.add_paragraph()
        _rtl_paragraph_marks(p)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(14)
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '18')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    # Header block: larger type + air before the first rule.
    _rtl_paragraph(
        f"הדיווחים של {student}",
        size_pt=30,
        bold=True,
        space_before=0,
        space_after=8,
    )
    _rtl_paragraph(
        f"({pronoun_dict['mentor']} - {teacher_line})",
        size_pt=18,
        bold=False,
        space_before=0,
        space_after=6,
    )
    _rtl_paragraph(
        f"שנת {_historic_hebrew_year_label(hebrew_year)}, "
        f"{_HISTORIC_TRIMESTER_HEBREW_LABELS[trimester_num]}",
        size_pt=18,
        bold=True,
        space_before=0,
        space_after=6,
    )
    _rtl_paragraph(
        f"דיווח היסטורי - הופק {get_printable_date()}",
        size_pt=13,
        bold=False,
        space_before=0,
        space_after=18,
    )

    _horizontal_rule()

    for idx, evaluation in enumerate(completed):
        klass = evaluation.evaluated_class
        header = f"{klass.name} - {klass.teacher}"
        if not evaluation.is_student_in_class:
            header += f" ({evaluation.student.first_name} עזב/ה את השיעור)"
        _rtl_paragraph(header, size_pt=12, bold=True, space_before=12 if idx else 0, space_after=6)

        text = (evaluation.evaluation_text or "").strip()
        if text:
            for line in text.splitlines():
                _rtl_paragraph(
                    line if line else " ",
                    size_pt=11,
                    bold=False,
                    space_before=0,
                    space_after=4,
                )
        else:
            _rtl_paragraph(" ", size_pt=11, space_before=0, space_after=4)

        _horizontal_rule()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


@_historic_staff_required
def historic_download_docx(request, student_id, hebrew_year, trimester_num):
    try:
        student = Student.objects.get(id=student_id)
    except Student.DoesNotExist:
        return render(request, 'common/general_error_page.html',
                      {'error_message': 'תלמיד/ה זה/זו לא נמצא/ת בבית הספר'})

    if trimester_num not in _HISTORIC_TRIMESTER_HEBREW_LABELS:
        return render(request, 'common/general_error_page.html',
                      {'error_message': 'סמסטר לא תקין'})

    data = _historic_build_evaluations_docx(student, hebrew_year, trimester_num)
    response = HttpResponse(
        data,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    filename = (
        f"דיווחים היסטוריים של {student} - "
        f"שנת {_historic_hebrew_year_label(hebrew_year)}, פגישה {trimester_num}.docx"
    )
    response['Content-Disposition'] = f"attachment; filename*=UTF-8''{quote(filename)}"
    return response
